from django.shortcuts import get_object_or_404, render,redirect
from django.template import context
from django.http import HttpResponse
from django.http import HttpResponseRedirect
from django.urls import reverse

from django.template.loader import get_template
from io import BytesIO
from xhtml2pdf import pisa
import datetime
from django.utils import timezone
from django.db.models import Count,Q
from django.db.models import Sum
from django.core.paginator import EmptyPage,PageNotAnInteger,Paginator

from accounts.models import Account
from django.contrib.auth.decorators import login_required
from django.core.paginator import EmptyPage,PageNotAnInteger,Paginator
from .models import Banner
from django.contrib.auth import authenticate,login,logout
from django.utils.text import slugify
from store.models import Product
from django.contrib import messages
from django.views.generic import View
from category.models import category
from orders.models import *
from docx import Document

from django.http import JsonResponse
from offers.models import Coupon,product_offer,cat_offer
from django.views.decorators.cache import never_cache
# Create your views here.

@never_cache
def admin_login(request):
    if request.method == 'POST':
        email=request.POST.get('adminname')
        password=request.POST.get('adminpassword')
        user=authenticate(email=email,password=password)
        # print(user.is_admin)
        if user is not None:
            login(request,user)
            if request.user.is_superuser == True:
                return redirect('admin_home')
            else:
                messages.error(request,'Your email or password is Incorrect!!')
                return redirect('admin_login')
        else:
            messages.error(request,'You are not an admin!!!')
            return redirect('admin_login') 

    return render(request,'admin/admin-login.html')

    
        
@never_cache
@login_required(login_url='admin_login')
def admin_home(request):
    if request.user.is_superuser==True:
       
       print('haiiiii')
    #    return render(request,'admin/admin-home.html')
       total_users=Account.objects.filter(blocked=False).count()
       total_products=Product.objects.filter(is_available=True).count()
       total_orders=Order.objects.filter(status='Delivered').count()
       total_revenue=Order.objects.filter(status='Delivered').aggregate(Sum('order_total'))
       
        # daily sales
       current_year=timezone.now().year
       print("current year : ",current_year)
       order_details=Order.objects.filter(created_at__lt=datetime.date(current_year,12,31),status="Delivered")
       monthly_order_count=[]
       month=timezone.now().month
       for i in range(1,month+2):
           monthly_order=order_details.filter(created_at__month=i).count()
           monthly_order_count.append(monthly_order)

        #    monthly sales
       today=datetime.datetime.now()
       dates=Order.objects.filter(created_at__month=today.month).values('created_at__day').annotate(order_items=Count('id')).order_by('created_at__day')
       returns=Order.objects.filter(created_at__month=today.month).values('created_at__day').annotate(retunrs=Count('id',filter=Q(status='Cancelled'))).order_by('created_at__day')
       sales=Order.objects.filter(created_at__month=today.month).values('created_at__day').annotate(sales=Count('id',filter=Q(status='Delivered'))).order_by('created_at__day')
       

    #    Most moving product
       most_moving_product_count=[]
       most_moving_product=[]
       products=Product.objects.all()
       for i in products:
           most_moving_product.append(i)
           most_moving_product_count.append(OrderProduct.objects.filter(product=i).count())

           #  status count
       placed_count=Order.objects.filter(status='placed').count()
       shipped_count=Order.objects.filter(status='Shipped').count()
       delivered_count=Order.objects.filter(status='Delivered').count()
       return_count=OrderProduct.objects.filter(status='Returned').count()
       cancelled_count=Order.objects.filter(status='Cancelled').count()
       print("total_orders",total_orders)
       print("total_products",total_products)
       print("total_users",total_users)
       print("total_revenue",total_revenue)
       print("monthly_order_count",monthly_order_count)
       print("today",today)
       print("dates",today)
       
       return render(request,'admin/admin-home.html',{
       
            'total_orders':total_orders,
            'total_products':total_products,
            'total_users':total_users,
            'total_revenue':total_revenue,
            'monthly_order_count': monthly_order_count,
            'today':today,
            'sales':sales,
            'returns':returns,
            'dates':dates,
            'most_moving_product':most_moving_product,
            'most_moving_product_count':most_moving_product_count,
            'status_count':[
                placed_count,
                shipped_count,
                delivered_count,
                return_count,
                cancelled_count,
             ]
        })
           

        
    return redirect('admin_login')


@never_cache
def admin_logout(request):
    logout(request)
    return redirect('admin_login')

# for seeing banners 
def banners(request):
    ban=Banner.objects.all()
    paginator=Paginator(ban,3)
    page=request.GET.get('page')
    paged_new=paginator.get_page(page)
    return render(request,'admin/Banners.html',{'paged_new':paged_new})

def addBanner(request):

    if request. method == 'POST':
        image=request.FILES.get('Image')
        ban=Banner()
        ban.images=image
        ban.save()
        return redirect('banners')
    else:
        return render(request,'admin/addBanner.html')
    
def editBanner(request,id):
  
    if request.method == 'POST':
        print("edit banner")

        ban=Banner.objects.get(id=id)
        ban.images=request.FILES.get('Image')
        ban.save()
        print("post: ",id)
        return redirect('banners')
    else:
        ban=Banner.objects.get(id=id)
        print(id)
        return render(request,'admin/editBanner.html',{'ban':ban})

def removeBanner(request,id) :
    ban=Banner.objects.get(id=id)
    ban.delete()
    return redirect('banners')



#for seeing userlist for admin
@login_required(login_url='admin_login')
def admin_userlist(request):
    print('>>>>>>>>>>')
    new=Account.objects.filter(is_superuser=False).order_by('id')
    paginator=Paginator(new,4)
    page_number=request.GET.get('page')

    if page_number is None:
        page_number = 1
    else:
        page_number = int(page_number)

    paged_new=paginator.get_page(page_number)
    #context={'userlist':new}
    #print(context)
    if page_number > paginator.num_pages:
        return HttpResponseRedirect(reverse('admin_userlist')+'?page=' + str(paginator.num_pages))
    print(paged_new)
    return render(request,'admin/admin-user.html',{'new':paged_new})

#for blocking user
def admin_block(request,id):
    print('blocked')
    user=Account.objects.get(id=id)
    user.blocked=True
    user.save()
    print(user.blocked)
    return redirect('admin_userlist')


#for unblocking user
def admin_unblock(request,id):
    print('unblocked')
    users=Account.objects.get(id=id)
    users.blocked=False
    users.save()
    print(users)
    return redirect('admin_userlist')

#to see the categorylist
@login_required(login_url='admin_login')
def admin_category(request):
    print('}}}}}}}}}}}}}}')
    new=category.objects.all()
    
    return render(request,'admin/admin-category.html',{'new':new})

#adding category
@login_required(login_url='admin_login')
def admin_addcategory(request):
    if request.method == 'GET':
       return render(request,'admin/admin-addcategory.html')
    else:
        if request.method == 'POST':
           print('post')
           name=request.POST['catname']
           description=request.POST['catdescription']
           image=request.FILES['Image']
           print("Imageeeee",image)
           
           cat=category()
           cat.category_name=name
           cat.cat_slug=slugify(name)
           cat.description=description
          
           cat.category_image=request.FILES['Image']
           cat.save()
           print("fileeeee: ",request.FILES)
           return redirect('admin_category')

def delete_category(request):
    id=request.GET.get('id')
    remove=category.objects.get(id=id).delete()
    return JsonResponse({'status':'true'})

def edit_category(request,id):
    cat=category.objects.get(id=id)
    return render(request,'admin/admin-editcategory.html',{'cat':cat})

#saving edited product to databse
@login_required(login_url='admin_login')
def update_category(request,id):
    if request.method == 'POST':
       Category=category.objects.get(id=id)
       Category_name=request.POST.get('catname')
       description=request.POST.get('catedescription')
       image=request.FILES.get('Image')
       print("REQUEST: ",request.FILES)
       if image:
           Category.category_image=image
       if Category_name:
            Category.category_name=Category_name
       if description:
         Category.description=description
       Category.save()
       messages.info(request,"Category updated")
    return redirect('admin_category')

# for seeing productlist
@login_required(login_url='admin_login')
def admin_productlist(request):
    product=Product.objects.all().order_by('id')
    paginator=Paginator(product,10)
    page=request.GET.get('page')
    paged_products=paginator.get_page(page)
    return render(request,'admin/admin-productlist.html',{'product':paged_products})

def prodsearch(request):
    if 's' in request.GET:
        keyword=request.GET['s']
        if keyword:
            products=Product.objects.order_by('id').filter(product_name__icontains=keyword)
        else:
            return HttpResponseRedirect(request.META["HTTP_REFERER"])
    paginator=Paginator(products,4)
    page=request.GET.get('page')
    paged_products=paginator.get_page(page)   
    context={
        'product':paged_products,
    }
    return render(request,'admin/admin-productlist.html',context)

@login_required(login_url='admin_login')
def admin_addproduct(request):
    if request.method == 'GET':
       cat=category.objects.all()
       
       return render(request,'admin/admin-product.html',{'cat':cat})
    
    else:
        product_name=request.POST['productname']
        description=request.POST['description']
        price=request.POST['price']
        # image=request.FILES['Image']
        cat=request.POST['category']
        print(cat,'//////////////')
        stock=request.POST['stock']
        slug=slugify(product_name)
        slug_compare=Product.objects.filter(prod_slug=slug)
        if slug == slug_compare:
            return redirect('admin_add')
        product=Product() 
        product.prod_slug=slugify(product_name)
        product.product_name=product_name
        product.description=description
        product.price=price
        product.user_price=price
        product.category_id=cat
        product.stock=stock
        if 'Image' in request.FILES:
              product.images=request.FILES['Image']
        if 'Image1' in request.FILES:
              product.image1=request.FILES['Image1']
        if 'Image2' in request.FILES:
              product.image2=request.FILES['Image2']
        product.save()
        if int(product.stock) < 1:
            product.is_available =False
        else:
            product.is_available = True
        product.save()

        return redirect('admin_productlist')



@login_required(login_url='admin_login')
def edit_product(request,id):
    if request.method == 'GET':
       prod=Product.objects.get(id=id)
       cat=category.objects.all()
       context={
        'prod':prod,
        'cat':cat
       }
       return render(request,'admin/admin-editproduct.html',context)
    else:
        product=Product.objects.get(id=id)
        product_name=request.POST.get('productname')
        description=request.POST.get('description')
        price=request.POST.get('price')     
        images=request.FILES.get('Image')
        image1=request.FILES.get('Image1')
        image2=request.FILES.get('Image2')
        cat=request.POST.get('category')
        stock=request.POST.get('stock')
        if product_name:
            product.product_name=product_name
        if description:
            product.description=description
        if price:
            product.price=price
            product.user_price=price

        if images:
            product.images=images

        if image1:
            product.image1=image1

        if image2:
            product.image2=image2

        product.stock=stock
        product.save()
        return redirect('admin_productlist')


def remove_product(request,id):
    remove=Product.objects.get(id=id).delete() 
    return redirect('admin_productlist')

# for adding product offer
def addproduct_Offer(request,id):
    if request.method == 'POST':
        prod_offer=product_offer()
        prod_offer.name=request.POST.get('productofferName')
        prod_offer.offer=request.POST.get('offer')
        prod_offer.offer_type=request.POST.get('offerType')
    
        prod_offer.product_id=id
        prod_offer.save()
        offer_price=(prod_offer.product.price)-(prod_offer.product.price*(int(prod_offer.offer)/100))
        prod_offer.product.p_offer=offer_price
        print(offer_price)
        print("product offer kandu pidichu kutta ")
        print(prod_offer.product.p_offer)
        prod_offer.product.save()

        if prod_offer.product.p_offer < 1 and prod_offer.product.cat_offer < 1:
            prod_offer.product.user_price = prod_offer.product.price

        elif prod_offer.product.p_offer < prod_offer.product.cat_offer and prod_offer.product.p_offer > 1:
            prod_offer.product.user_price = prod_offer.product.p_offer
        
        elif prod_offer.product.cat_offer < prod_offer.product.p_offer and prod_offer.product.cat_offer > 1:
            prod_offer.product.user_price = prod_offer.product.cat_offer

        elif prod_offer.product.cat_offer > 1 and prod_offer.product.p_offer == 0:
            prod_offer.product.user_price =prod_offer.product.cat_offer
        
        elif prod_offer.product.p_offer > 1 and prod_offer.product.cat_offer == 0:
            prod_offer.product.user_price = prod_offer.product.p_offer
        
        else:
            prod_offer.product.user_price = prod_offer.product.cat_offer
        
        prod_offer.product.save()


        return redirect('admin_productlist')

    else:
        product=Product.objects.get(id=id)
        return render(request,'admin/admin-prodOffer.html',{'product':product})
    
# for editing product Offer
def edit_prodOffer(request,id):
    
    if request.method =='POST':
        print(id)
        prod_offer=product_offer.objects.get(id=id)
        print("editttttttt")
        print(prod_offer)
        prod_offer.name=request.POST.get('productofferName')
        prod_offer.offer=request.POST.get('offer')
        prod_offer.offer_type=request.POST.get('offerType')
    
        # prod_offer.product_id=id
        prod_offer.save()
        offer_price=(prod_offer.product.price)-(prod_offer.product.price*(int(prod_offer.offer)/100))
        prod_offer.product.p_offer=offer_price
        print(offer_price)
        print("product offer kandu pidichu kutta ")
        print(prod_offer.product.p_offer)
        prod_offer.product.save()

        if prod_offer.product.p_offer < 1 and prod_offer.product.cat_offer < 1:
            prod_offer.product.user_price = prod_offer.product.price

        elif prod_offer.product.p_offer < prod_offer.product.cat_offer and prod_offer.product.p_offer > 1:
            prod_offer.product.user_price = prod_offer.product.p_offer
        
        elif prod_offer.product.cat_offer < prod_offer.product.p_offer and prod_offer.product.cat_offer > 1:
            prod_offer.product.user_price = prod_offer.product.cat_offer

        elif prod_offer.product.cat_offer > 1 and prod_offer.product.p_offer == 0:
            prod_offer.product.user_price =prod_offer.product.cat_offer
        
        elif prod_offer.product.p_offer > 1 and prod_offer.product.cat_offer == 0:
            prod_offer.product.user_price = prod_offer.product.p_offer
        
        else:
            prod_offer.product.user_price = prod_offer.product.cat_offer
        
        prod_offer.product.save()


        return redirect('admin_productlist')

    else:
        print(id)
        offer=product_offer.objects.get(product_id=id)
        print(offer)
        return render(request,'admin/edit_productOffer.html',{'offer':offer})


# removing prodOffer
def remove_productOffer(request):
    print(">>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>")
    print("keriiii")
    print(">>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>")
    id=request.GET.get('id')
    print(id)
    prod_offer=product_offer.objects.get(product_id=id)
    prod_offer.product.p_offer = 0
    prod_offer.product.save()
    
    if prod_offer.product.p_offer < 1 and prod_offer.product.cat_offer < 1:
        prod_offer.product.user_price = prod_offer.product.price

    elif prod_offer.product.p_offer < prod_offer.product.cat_offer and prod_offer.product.p_offer > 1:
        prod_offer.product.user_price = prod_offer.product.p_offer
        
    elif prod_offer.product.cat_offer < prod_offer.product.p_offer and prod_offer.product.cat_offer > 1:
        prod_offer.product.user_price = prod_offer.product.cat_offer

    elif prod_offer.product.cat_offer > 1 and prod_offer.product.p_offer == 0:
        prod_offer.product.user_price =prod_offer.product.cat_offer
    
    elif prod_offer.product.p_offer > 1 and prod_offer.product.cat_offer == 0:
        prod_offer.product.user_price = prod_offer.product.p_offer
        
    else:
        prod_offer.product.user_price = prod_offer.product.cat_offer
        
    prod_offer.product.save()
    prod_offer.delete()

    return JsonResponse({'status':'true'})


def add_catOffer(request,id):
    if request.method =='POST':
        catedory_offer=cat_offer()
        catedory_offer.name=request.POST.get('productofferName')
        catedory_offer.offer=request.POST.get('offer')
        catedory_offer.offer_type=request.POST.get('offerType')
      
        catedory_offer.category_id=id
        catedory_offer.save()
        print("category offer save aayito")
        catedory_offer.category.c_offer=True
        catedory_offer.category.save()

        prod=Product.objects.filter(category_id=id)

        for i in prod:
            i.cat_offer=i.price-(i.price*(int(catedory_offer.offer)/100))
            i.save()

            if i.p_offer < 1 and i.cat_offer < 1:
                i.user_price = i.price

            elif i.p_offer < i.cat_offer and i.p_offer > 1:
                i.user_price = i.p_offer
        
            elif i.cat_offer < i.p_offer and i.cat_offer > 1:
                i.user_price = i.cat_offer

            elif i.cat_offer > 1 and i.p_offer == 0:
                i.user_price =i.cat_offer
        
            elif i.p_offer > 1 and i.cat_offer == 0:
                i.user_price = i.p_offer
        
            else:
                i.user_price = i.cat_offer
                print("category offer calculate aaki")
            i.save()
        return redirect('admin_category')

    else:
        Category=category.objects.get(id=id)
        return render(request,'admin/admin-CatOffer.html',{'Category':Category})

# for editing cat offer
def edit_catOffer(request,id):
    if request.method =='POST':
        catedory_offer=cat_offer.objects.get(id=id)
        catedory_offer.name=request.POST.get('productofferName')
        catedory_offer.offer=request.POST.get('offer')
        catedory_offer.offer_type=request.POST.get('offerType')
      
        
        catedory_offer.save()
        c_id=catedory_offer.category_id
        print("iddddddddddddddddd",c_id)
        print("category offer save aayito")
        catedory_offer.category.c_offer=True
        catedory_offer.category.save()

        prod=Product.objects.filter(category_id=c_id)
        print("forrrrrrrrrrrrrrrrr")
        print(prod)
        for i in prod:
            print("proddddddddddddddddddd")
            i.cat_offer=i.price-(i.price*(int(float(catedory_offer.offer))/100))
            i.save()

            if i.p_offer < 1 and i.cat_offer < 1:
                i.user_price = i.price

            elif i.p_offer < i.cat_offer and i.p_offer > 1:
                i.user_price = i.p_offer
        
            elif i.cat_offer < i.p_offer and i.cat_offer > 1:
                i.user_price = i.cat_offer

            elif i.cat_offer > 1 and i.p_offer == 0:
                i.user_price =i.cat_offer
        
            elif i.p_offer > 1 and i.cat_offer == 0:
                i.user_price = i.p_offer
        
            else:
                i.user_price = i.cat_offer
                print("category offer calculate aaki")
            i.save()
        return redirect('admin_category')
        

    else:
        cate_offer=cat_offer.objects.get(category_id=id)
        return render(request,'admin/editCatOffer.html',{'cate_offer':cate_offer})


# for removing cat offer
def remove_catOffer(request):
    id=request.GET.get('id')
    print("idd : ",id)
    print("############################################################################################################")
    cate_offer=Product.objects.filter(category_id=id)
    print("cattttttttttt",cate_offer)
    print("******************************************************************************************************************")
    if cate_offer :
            for cat in cate_offer:
                cat.cat_offer = 0
                cat.save()
                print("???????????????????")
                print(cat.cat_offer)
                if cat.p_offer < 1 and cat.cat_offer < 1:
                    cat.user_price = cat.price

                elif cat.p_offer < cat.cat_offer and cat.p_offer > 1:
                    cat.user_price = cat.p_offer
                    
                elif cat.cat_offer < cat.p_offer and cat.cat_offer > 1:
                    cat.user_price = cat.cat_offer

                elif cat.cat_offer > 1 and cat.p_offer == 0:
                    cat.user_price =cat.cat_offer
                
                elif cat.p_offer > 1 and cat.cat_offer == 0:
                    cat.user_price = cat.p_offer
                    
                else:
                    cat.user_price = cat.cat_offer
                
                cat.save()
                print("okke calculate aaki")
            cater=cat_offer.objects.get(category_id=id)
            cater.category.c_offer=False
            cater.category.save()
            cater.delete()
    else:   
        cat=cat_offer.objects.get(category_id=id)
        cat.category.c_offer=False
        cat.category.save()
        
        cat.delete()

    return JsonResponse({'status':'true'})


#for seeing coupons 
def admin_coupons(request):
    coupon=Coupon.objects.all()
    return render(request,'admin/admin-coupons.html',{'coupon':coupon})

def admin_addcoupon(request):
    if request.method == 'POST':
       code=request.POST.get('code')
       if Coupon.objects.filter(coupon_code=code).exists():
           messages.error(request,'This code exist!!!')
           return redirect('admin_addcoupon')
       coupon=Coupon()
       coupon.coupon_name= request.POST.get('couponName')
       coupon.coupon_code = request.POST.get('code')
       coupon. discount=request.POST.get('discount')
       coupon.limited_price=request.POST.get('limit')
       coupon.valid_from=request.POST.get('validFrom')
       coupon.valid_to=request.POST.get('validTo')
       coupon.save()
       return redirect('admin_coupons')

    else:    
        return render(request,'admin/admin_addcoupon.html')
    
def RemoveCoupon(request,id):
    coupon=get_object_or_404(Coupon,id=id)
    coupon.delete()
    return redirect('admin_coupons')


# for seeing order list
def admin_orderlist(request):
    print("orders")
    orderlist = Order.objects.all().order_by('-id')
    print(orderlist)
    paginator=Paginator(orderlist,10)
    page=request.GET.get('page')
    paged_orderlist=paginator.get_page(page)

    return render(request,'admin/orderList.html',{'orderlist':paged_orderlist} )

def viewDetails(request,id):
    orders=OrderProduct.objects.filter(order=id)
    print(orders)
    return render(request,'admin/admin_VewDetails.html',{'orders':orders})


# for changing  order status
def editStatus(request):
    id=request.GET.get('id')
    stat=request.GET.get('status')
    print("iddddddd:",id)
    print("status :",stat)

    order=Order.objects.get(id=id)
    orderItem=OrderProduct.objects.filter(order_id=id)
    print("ordersssssss:",orderItem)
    print(order)
    if stat == 'Cancel':
        order.status='Cancelled'
        for i in orderItem:
            i.status='Cancelled'
            i.save()
            print(i.status)
    elif stat == 'Shipped':
        order.status = 'Shipped'
        # for i in orderItem:
            # i.status = 'Shipped'
    elif stat == 'Delivered':
        order.status = 'Delivered'
    
    print("ddddddddddddd")
    order.save()
    print(order.status)
    return JsonResponse({})

# Sales Report
def salesReport(request):
    orders=Order.objects.filter(status='Delivered').order_by('id')
    new_order_list=[]
    for i in orders:
        order_item=OrderProduct.objects.filter(order_id=i.id)
        for j in order_item:
            item={
                'id':i.id,
                'ordered_date':i.created_at,
                'price':i.order_total,
                'method':i.pay_method,
                'user':j.user.username,
                'status':i.status,
            }
            new_order_list.append(item)
    paginator=Paginator(new_order_list,10)
    page=request.GET.get('page')
    paged_orderlist=paginator.get_page(page)
    return render(request,'admin/salesReport.html',{'order':paged_orderlist})

# sales report by date

def by_date(request):
    if request.GET.get('from'):
        sales_date_from=datetime.datetime.strptime(request.GET.get('from'),"%Y-%m-%d")
        sales_date_to=datetime.datetime.strptime(request.GET.get('to'),"%Y-%m-%d")
        sales_date_to+=datetime.timedelta(days=1)
        orders=Order.objects.filter(created_at__range=[sales_date_from,sales_date_to],status='Delivered').order_by('id')

        new_order_list=[]
        for i in orders:
            order_items=OrderProduct.objects.filter(order_id=i.id)
            for j in order_items:
                item={
                'id':i.id,
                'ordered_date':i.created_at,
                'price':i.order_total,
                'method':i.pay_method,
                'user':j.user.username,
                'status':i.status,
                }
                new_order_list.append(item)
    else:
        messages.error(request,'Select fields before submiting!!!')
        return redirect('salesReport')
    
    return render(request,'admin/salesReport.html',{'order':new_order_list})

# sales report by 
def by_month(request):
    month=request.GET.get('month')
    orders=Order.objects.filter(created_at__month=month,status='Delivered').order_by('id')
    print(orders)
    new_order_list=[]
    for i in orders:
        order_items=OrderProduct.objects.filter(order_id=i.id)
        for j in order_items:
            item={
                'id':i.id,
                'ordered_date':i.created_at,
                'price':i.order_total,
                'method':i.pay_method,
                'user':j.user.username,
                'status':i.status,
            }
            new_order_list.append(item)
    return render(request,'admin/salesReport.html',{'order':new_order_list})


# for generates sales report
class genereateSalesreport(View):
    def get(self,request,*args,**kwargs):
        try:
            orders=Order.objects.all()
            new_order_list=[]
            print("sales  Report")
            print("ordersssssss:",orders)
            for i in orders:
                order_item=OrderProduct.objects.filter(order_id=i.id)
                print("order_items",order_item)
                for j in order_item:
                    print("sales Ordersssssss")
                    item={
                        'id':i.id,
                        'ordered_date':i.created_at,
                        'price':i.order_total,
                        'method':i.pay_method,
                        'user':j.user.username,
                        'status':i.status,

                        }
                    new_order_list.append(item)
        except:
            return HttpResponse("505 not found")
        data={
            'order':new_order_list
        }
        pdf=render_to_pdf('admin/salesReport_pdf.html',data)
        return HttpResponse(pdf,content_type='application/pdf')

def download_docx(request):
    document = Document()
    document.add_heading('Sales Report',0)
    records = (
        (3, '101', 'Spam'),
        (7, '422', 'Eggs'),
        (4, '631', 'Spam')
    )

    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Qty'
    hdr_cells[1].text = 'Id'
    hdr_cells[2].text = 'Desc'

    for qty, id, desc in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(qty)
        row_cells[1].text = id
        row_cells[2].text = desc

    document.add_page_break()

    response = HttpResponse(content_type = 'application/vnd.openxmlformats-officedocuments.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=sales-report.docx'
    document.save(response)

    return response


def render_to_pdf(template_src,context_dict={}):
    template=get_template(template_src)
    html=template.render(context_dict)
    result=BytesIO()
    pdf=pisa.pisaDocument(BytesIO(html.encode("ISO-8859-1")),result)
    if not pdf.err:
        return HttpResponse(result.getvalue(),content_type='application/pdf')
    return None
